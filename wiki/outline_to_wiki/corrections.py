from docx import Document
from docx.shared import Pt

from vars import doc, paragraphs

a = 0
b = str()
c = str()
#Correction one runs if two former paragraphs need to formated into bullet points
def correction():
    for p in doc.paragraphs:
        if p.text.startswith('TELLZ'):
            if p.text.startswith(b) or p.text.startswith(c):
                p.text = p.text.replace('TELLZ ', '*<h3>')
                p.text = p.text.replace(' SELLZ', '</h3>')

#Correction Two Removes Custom strings from non bullet text
def correction_two():
    for p in doc.paragraphs:
        if p.text.startswith('TELLZ'):
            if p.text.startswith(b) or p.text.startswith(c):
                p.text = p.text.replace('TELLZ ', '')
                p.text = p.text.replace(' SELLZ', '')

def test(a,b,c):
    for p in doc.paragraphs: 
        if p.style.name.startswith('Heading 2') and p.text.startswith('<') == False and p.text.startswith('*') == False:
            '''
            Asigns random string to beginig and end of text.
            This lets us go back later and replace them.
            Will replace Text with random string of numbersand letters later 
            so that they dont get mixed up with real words.
            '''
            a = a + 1
            p.text = 'TELLZ ' + p.text + ' SELLZ'
            #Stores text to refrence later during correction
            if a < 3:
                if a == 1:
                    b = p.text

                elif a == 2:
                    c = p.text
            
            #Changes third or more paragraphs into bullets
            elif a >= 3 and p.style.name.startswith('Heading 2') and p.text.startswith('TELLZ '):
                p.text = p.text.replace('TELLZ ', '*<h3>')
                p.text = p.text.replace(' SELLZ', '</h3>')

        #Makes corrections and resets variables whenever heading is no longer heading 2
        elif p.style.name.startswith('Heading 2') == False:
            if a < 3:
                correction_two()
                a = 0
                b = str(' ')
                c = str(' ')

            elif a >= 3:
                correction()
                a = 0
                b = str(' ')
                c = str(' ')

