'''
This  is where i test and debug functions that aren't working
A good method is to print the paragraph or run's name and go from there

'''
#NO TOUCH vvvvvv
from docx import Document
from docx.shared import Pt

doc_path = r'C:\Users\Cliff\Desktop\RE-SEED.docx'
doc = Document(doc_path)
paragraphs = doc.paragraphs
sections = doc.sections
#NO TOUCH ^^^^^

''' Examples

#Checks runs
def bold_text():
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                run.text = '<b>' + run.text + '</b>'
                run.bold = False

#Checks para
for paragraph in paragraphs:
        if paragraph.style.name.startswith('Title'):
            paragraph.text = '<center><h1>' + paragraph.text + '</h1></center>'
            paragraph.style = doc.styles['Normal']if paragraph.style.name.startswith('Heading 2'):
            i = i + 1
        else:
            continue

        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading 2') and i >= 3:
                paragraph.text = '* <h3>' + paragraph.text + '</h3>'
                paragraph.style = doc.styles['Normal']
                if paragraph.style.name.startswith('Heading 2') == False:
                    i = 0

            elif paragraph.style.name.startswith('Heading 2') and i < 3:
                paragraph.text = '<h3>' + paragraph.text + '</h3>'
                paragraph.style = doc.styles['Normal']
                if paragraph.style.name.startswith('Heading 2') == False:
                    i = 0
    print(paragraph.text)
    
    print(aa)if d == True:
            if p.text.startswith(b):
                p.text = p.text.replace('one', '*<h3>')
                p.text = p.text.replace('two', '</h3>')
            if p.text.startswith(c):
                p.text = p.text.replace('one', '*<h3>')
                p.text = p.text.replace('two', '</h3>')
                d = False  
'''

#NO TOUCH
#Doc file            C:\Users\Cliff\Desktop\test.docx
#My Outline          C:\Users\Cliff\Desktop\May 16th 2024 - Copy.docx
#Michael Outline    C:\Users\Cliff\Desktop\RE-SEED.docx

#account for spaces i.e starts with ' '
#WIP NEED TO ACCOUNT FOR FIRST TWO PARA
#could save a val in var and then go back to refrence it
a = 0
b = str()
c = str()
d = bool()
e = 0

def correction():
    for p in doc.paragraphs:
        if p.text.startswith('<'):
            if p.text.startswith(b) or p.text.startswith(c) :
                p.text = '*' + p.text

def test(a,b,c,d):
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading 2') and p.text.startswith('<h3>') == False and p.text.startswith('*<h3>') == False:
            p.text = 'one' + p.text + 'two'

            if a < 3:
                p.text = p.text.replace('one', '<h3>')
                p.text = p.text.replace('two', '</h3>')

                if a == 1:
                    b = p.text

                elif a == 2:
                    c = p.text

            elif a >= 3 and p.text.startswith:
                p.text = p.text.replace('one', '*<h3>')
                p.text = p.text.replace('two', '</h3>')
                
                correction()
            a = a + 1
        elif p.style.name.startswith('Heading 2') == False:
            a = 0
            b = str()
            c = str()

test(a,b,c,d)

for p in doc.paragraphs:
    print(p.text)



