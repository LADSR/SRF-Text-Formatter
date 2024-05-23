from docx import Document
from docx.shared import Pt

from vars import doc, paragraphs

def bold_text():
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                run.text = '<b>' + run.text + '</b>'
                run.bold = False

def italics_text():
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.italic:
                run.text = '<i>' + run.text + '</i>'
                run.italic = False

def underline_text():
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.underline:
                run.text = '<i>' + run.text + '</i>'
                run.underline = False

def strike_text():
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.strike:
                run.text = '<s>' + run.text + '</s>'
                run.font.strike = False

def strike_text():
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.strike:
                run.text = '<s>' + run.text + '</s>'
                run.font.strike = False

def sub_text():
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.subscript:
                run.text = '<SUB>' + run.text + '</SUB>'
                run.font.subscript = False

def sup_text():
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.superscript:
                run.text = '<SUP>' + run.text + '</SUP>'
                run.font.superscript = False

def other_text_styles():
    bold_text()
    italics_text()
    underline_text()
    strike_text()
    sup_text()
    sub_text()
