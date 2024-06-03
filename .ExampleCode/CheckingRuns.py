'''This code takes bold text and formats it into mediawiki syntax'''

from docx import Document
from docx.shared import Pt

doc_path = r'PATH HERE'
doc = Document(doc_path)
paragraphs = doc.paragraphs

def bold_text():
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                run.text = '<b>' + run.text + '</b>'
                run.bold = False
    print(paragraph.text)