'''
This is where i test and debug functions that aren't working
A good method is to print the paragraph or run's text and go from there
'''
from docx import Document
from docx.shared import Pt
path = input('Select doc: ')

def h2_correction_one(a, b, c):
    for paragraph in doc.paragraphs:
            if paragraph.text.startswith('tOdjewZ7wd '):
                if paragraph.text.startswith(b) or paragraph.text.startswith(c):
                    paragraph.text = paragraph.text.replace('tOdjewZ7wd ', '')
                    paragraph.text = paragraph.text.replace(' tUDK9aIVL1', '')

#Correction Two Removes Custom strings from non bullet text
def h2_correction_two(a, b, c):
    from vars import doc, paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith('tOdjewZ7wd '):
            if paragraph.text.startswith(b) or paragraph.text.startswith(c):
                paragraph.text = paragraph.text.replace('tOdjewZ7wd ', '*<h3>')
                paragraph.text = paragraph.text.replace(' tUDK9aIVL1', '</h3>')

def heading_2(a, b, c):
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading 2') and paragraph.text.startswith('<') == False and paragraph.text.startswith('*') == False and paragraph.text.startswith('tOdjewZ7wd') == False:
                a = a + 1
                paragraph.text = 'tOdjewZ7wd ' + paragraph.text + ' tUDK9aIVL1'
                #Stores text to refrence later during correction
                if a < 3:
                    if a == 1:
                        b = paragraph.text

                    elif a == 2:
                        c = paragraph.text
                
                #Changes third or more paragraphs into bullets
                elif a >= 3 and paragraph.style.name.startswith('Heading 2') and paragraph.text.startswith('tOdjewZ7wd'):
                    paragraph.text = paragraph.text.replace('tOdjewZ7wd ', '*')
                    paragraph.text = paragraph.text.replace(' tUDK9aIVL1', '')

        #Makes corrections and resets variables whenever heading is no longer heading 2
        elif paragraph.style.name.startswith('Heading 2') == False:
            if a < 3:
                h2_correction_one()
                a = 0
                b = str('')
                c = str('')

            elif a >= 3:
                h2_correction_two()
                a = 0
                b = str('')
                c = str('')